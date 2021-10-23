import os

main_dir = r"C:\Users\sugarkhuu\Documents\python\Python_Introduction\week4"
os.chdir(main_dir)


import glob

ext = 'pdf'
for filepath in glob.glob('*.' + ext):
    print(filepath)
    
glob.glob('*.pdf')
glob.glob('*.csv')
glob.glob('khan*.pdf')
glob.glob('*k*.pdf')



# EXCEL
import pandas as pd
xlsx = pd.ExcelFile('data1.xlsx')
xlsx.sheet_names
df = xlsx.parse("Sheet1")
df


# openpyxl
from openpyxl import load_workbook


# read
wb = load_workbook("data.xlsx")
print(wb.sheetnames)     
ws = wb["double"]        # activate sheet "double" as ws
print(ws.title)          # sheet name
print(ws["D10"].value)   # cell value by cell name
ws.cell(row=10, column=4).value # cell value by cell location
ws["A1:C2"][1][2].value         # cell value from selected range
ws["A:B"][1][1].value           # cell value from selected range
ws[1:2][1:3]                    # row and column range


# iter_rows, iter_cols
for value in ws.iter_rows(min_row=1,
                             max_row=2,
                             min_col=1,
                             max_col=3,
                             values_only=True): # values_only=True
    print(value)

for value in ws.iter_cols(min_row=1,
                             max_row=2,
                             min_col=1,
                             max_col=3,
                             values_only=True): # values_only=True
    print(value)

# rows, columns
for row in ws.rows:
    # print(row) # row[1].value
    print(row[1].value)


# restore and dump - database
# dumping to json
import json
products = {}

# Using the values_only because you want to return the cells' values
for row in ws.iter_rows(min_row=2,
                           min_col=1,
                           max_col=4,
                           values_only=True):
    product_id = row[0]
    product = {
        "firstName": row[1],
        "lastName": row[2],
        "citizenId": row[3]
    }
    products[product_id] = product

with open('data.json', 'w') as myfile:
    json.dump(products, myfile, sort_keys=True, indent=4) # 


# loop through cell    
for cell in ws[1]:
    print(cell.value)

# ws.max_row
# ws.max_column


# write
from openpyxl import Workbook
wb = Workbook() # instantiate a workbook
ws = wb.active  # select first sheet of wb as sheet

ws["A1"] = "11 sar"
ws["B1"] = "good"

wb.save(filename="hello_world.xlsx") 


wb = load_workbook(filename="hello_world.xlsx")
ws = wb['Sheet']


cell = ws["A1"]
cell.value = "hey"

wb.save(filename="hello_world.xlsx")

wb = Workbook() # instantiate a workbook
wb.save(filename="hello_world.xlsx")


wb = Workbook() # instantiate a workbook
fs = wb.create_sheet("firstSheet",index=None)
wb.save(filename="hello_world.xlsx")


wb = load_workbook(filename="hello_world.xlsx")
wb.remove(wb['firstSheet'])
wb.save(filename="hello_world.xlsx")


wb = load_workbook(filename="hello_world.xlsx")
wb.sheetnames
wb.copy_worksheet(wb['Sheet'])
ws = wb['Sheet Copy']
ws.title = 'diffSheet'
wb.save(filename="hello_world.xlsx")


wb = load_workbook(filename="data.xlsx")
wb.sheetnames
wb.copy_worksheet(wb['Sheet1'])
ws = wb['Sheet1 Copy']
ws.title = 'formula'
ws["E12"] = '=AVERAGE(E2:E11)'
wb.save(filename="data.xlsx")


wb = load_workbook(filename="data.xlsx")
wb.sheetnames
wb.copy_worksheet(wb['Sheet1'])
ws = wb['Sheet1 Copy']
ws.title = 'formula2'

last_row = ws.max_row
last_col = ws.max_column
age_col  = 5

ws.cell(last_row+1,age_col).value = '=AVERAGE(E2:E' + str(last_row) + ')'
wb.save(filename="data.xlsx")



# with pandas, multiple sheets


xl = pd.ExcelFile("data.xlsx")
df = xl.parse('Sheet1')
with pd.ExcelWriter('multiplesheet.xlsx', engine='xlsxwriter') as writer:
    df.to_excel(writer, sheet_name='Sheet1')
    df.to_excel(writer, sheet_name='Sheet2')
    df.to_excel(writer, sheet_name='Sheet3')
writer.save()



# writer = pd.ExcelWriter('example.xlsx', engine='xlsxwriter')
# df.to_excel(writer, 'Sheet1') # Multiple sheets



# Main source - https://realpython.com/openpyxl-excel-spreadsheets-python/
# https://xlsxwriter.readthedocs.io/
# https://www.datacamp.com/community/tutorials/python-excel-tutorial
# https://www.dataquest.io/blog/excel-and-pandas/
# https://www.python-excel.org/


# CSV

import csv
with open('data.csv', 'r') as file:
    reader = csv.reader(file)
    for row in reader:
        print(row)
        

with open('data.csv', 'a', newline='') as file:
    writer = csv.writer(file)
    writer.writerow(["11", "Gerel", "Erdene"])
    writer.writerow(["12", "Gan", "Tuguldur"])
    writer.writerow(["13", "Chimeg", "Bayar"])
    

csv_rowlist = [["14", "Saran", "Khurel"], ["15", "Amar", "Khand"],
               ["16", "Tuvshin", "Tulga"]]
with open('data.csv', 'a', newline='') as file:
    writer = csv.writer(file)
    writer.writerows(csv_rowlist)
    


with open('data_dict.csv', 'w', newline='') as file:
    fieldnames = ['id', 'firstName']
    writer = csv.DictWriter(file, fieldnames=fieldnames)

    writer.writeheader()
    writer.writerow({'id': '17', 'firstName': 'Balbar'})
    writer.writerow({'id': '18', 'firstName': 'Ochir'})
    
    


# pandas works


# PDF

import PyPDF2 # pip install PyPDF2

# read
anna = PyPDF2.PdfFileReader('anna_karenina.pdf')
nPages = anna.getNumPages()
docInfo  = anna.documentInfo
docInfo.author

nPages = anna.getNumPages()
page = anna.getPage(0)
page_content = page.extractText()
print(page_content.encode('utf-8'))


# write
from PyPDF2 import PdfFileWriter
pdf_writer = PdfFileWriter()

page = pdf_writer.addBlankPage(width=72, height=72)

# pdf_writer.write(str('output.pdf'))

with open('output.pdf','wb') as output_file:
    pdf_writer.write(output_file)



# slice
input_pdf = PyPDF2.PdfFileReader('anna_karenina_long.pdf')
pdf_writer = PyPDF2.PdfFileWriter()
for n in range(1, 4):
    page = input_pdf.getPage(n)
    pdf_writer.addPage(page)
    
with open('anna_2_4.pdf','wb') as output_file:
    pdf_writer.write(output_file)
   

   
input_pdf = PyPDF2.PdfFileReader('khanbank.pdf')
pdf_writer = PyPDF2.PdfFileWriter()
for n in range(1, 4):
    page = input_pdf.getPage(n)
    pdf_writer.addPage(page)
    
with open('khan_2_4.pdf','wb') as output_file:
    pdf_writer.write(output_file)




from PyPDF2 import PdfFileMerger
pdf_merger = PdfFileMerger()

pdfs = list(glob.glob("*.pdf"))

for path in pdfs:
    pdf_merger.append(str(path))

with open('all_pdf.pdf','wb') as output_file:
    pdf_merger.write(output_file)


import fpdf #pip3 intall fpdf

pdf = fpdf.FPDF(format='letter') #pdf format
pdf.add_page() #create new page
pdf.set_font("Arial", size=12) # font and textsize
pdf.cell(40, 2, txt="Happy families", ln=0, align="L")
pdf.cell(100, 2, txt="are all alike;", ln=2, align="L")
pdf.cell(200, 10, txt="every unhappy family", ln=2, align="L")
pdf.output("test1.pdf")

# pdf.cell(200, 10, txt="is unhappy", ln=0, align="R")
# pdf.cell(200, 10, txt="in its own way.", ln=0, align="L")
# pdf.cell(200, 10, txt="haha.", ln=1, align="L")

# ln 0,1,2

# pip install reportlab
from reportlab.lib.colors import blue
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
from reportlab.pdfgen.canvas import Canvas

canvas = Canvas("font-colors.pdf", pagesize=LETTER)
canvas.setFont("Times-Roman", 12) # Set font to Times New Roman with 12-point size
canvas.setFillColor(blue) # Draw blue text one inch from the left and ten inches from the bottom
canvas.drawString(1 * inch, 10 * inch, "Blue text")
canvas.save() # Save the PDF file



# Other
# import tabula # pip install tabule-py - Java should be installed
# tabular-py
# import textract
# pdfkit


# Word

# pip install python-docx
from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

document.add_picture('horse.png', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('demo.docx')



# convert PDF to Word
anna = PyPDF2.PdfFileReader('anna_karenina.pdf')
page = anna.getPage(0)
page_content = page.extractText()
document = Document()
document.add_paragraph(page_content)
document.save('anna_karenina.docx')  



# read
document = Document('anna_karenina.docx')
for para in document.paragraphs:
    print(para.text)



# Conclusion
# https://www.programiz.com/python-programming/methods/built-in/open


f = open("anna.txt")   

f = open("anna.txt")      # equivalent to 'r' or 'rt'
f = open("anna.txt",'w')  # write in text mode
f = open("anna.txt", mode='r', encoding='utf-8')
f.close()

try:
   f = open("anna.txt", encoding = 'utf-8')
   # perform file operations
finally:
   f.close()
   
   
with open("anna.txt",'w',encoding = 'utf-8') as f:
   f.write("my first file\n")
   f.write("This file\n\n")
   f.write("contains three lines\n")
   
   
f = open("anna.txt",'r',encoding = 'utf-8')
f.read(4)    # read the first 4 data
f.read(4)    # read the next 4 data
f.read()     # read in the rest till end of file
f.tell()    # get the current file position
f.seek(0)   # bring file cursor to initial position
for line in f:
    print(line, end = '')



# REcognize cyrillic Mongolian letters - utf-8-sig
df = pd.read_excel('cyrillic.xlsx', encode = 'utf-8-sig')