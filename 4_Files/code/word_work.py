#%% Working with WORD files

#%% environment
import os
import sys
from PyPDF2 import PdfFileReader

main_dir = r"D:\Documents\python\repo\Introduction_Python\4_Files"
sys.path.append(r'' + main_dir + '\code')
from settings import *

# set main dir
os.chdir(main_dir + '/' + resource_dir)

#%% write
# pip install python-docx
from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

document.add_picture('horse.png', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

tmp = """ Happy families are all alike; every unhappy family 
    is unhappy in its own way. Everything was in confusion 
    in the Oblonskys’ house. The wife had discovered that the 
    husband was carrying on an intrigue with a French girl, who 
    had been a governess in their family, and she had announced to 
    her husband that she could not go on living in the same house with 
    him. This position of affairs had now lasted three days, and not only 
    the husband and wife themselves, but all the members of their family and 
    household, were painfully conscious of it. """

import re
tmp = re.sub("\n","",tmp)


p = document.add_paragraph(tmp)
p.paragraph_format.left_indent = Inches(.25)

document.save('../' + result_dir + '/' + 'demo.docx')


# Style
https://stackoverflow.com/questions/27904470/checking-for-particular-style-using-python-docx

from docx import Document
d = Document('test.docx')
# for paragraph in d.paragraphs:
#     for run in paragraph.runs:
#         print (run.font.size)
#         print(run)

a = d.paragraphs    
a[0].runs[0].text
a[0].runs[0].bold = True
a[0].runs[0].text = 'Hello'


d = Document('../' + result_dir + '/' + 'demo.docx')
a = d.paragraphs    


# convert PDF to Word
anna = PdfFileReader('anna_karenina.pdf')
page = anna.getPage(0)
page_content = page.extractText()
page_content = page_content.replace('\n','')
document = Document()
document.add_paragraph(page_content)
document.save('../' + result_dir + '/' + 'anna_karenina.docx')  



# read
document = Document('../' + result_dir + '/' + 'anna_karenina.docx')
for i,para in enumerate(document.paragraphs):
    print(para.text)
    print(i)
# %%
